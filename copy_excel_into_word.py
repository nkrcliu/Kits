from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches
from docx.shared import RGBColor
import xlrd

document = Document()

wb = xlrd.open_workbook(r'1.xlsx')
sheetname = wb.sheet_names()[0]
sheet_c=wb.sheet_by_index(0)

for i in range(77):
    row_date = sheet_c.row_values(i)
    table = document.add_table(rows=3, cols=1)
    col = table.columns[0]
    col.width = Inches(0.5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '    '+row_date[0]+'    '+row_date[1]+'    '+row_date[2]
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = row_date[3]
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = row_date[4]
    #document.add_page_break()
    document.add_paragraph()

document.save('riji.docx')
